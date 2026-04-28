#!/usr/bin/env python3
"""10주차 주관식 문제지 v2 — DOCX 생성"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "/home/pc/dev/metade/10week/문제지/10주차_주관식_문제지_v2.docx"

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


# ── Style helpers ──
def add_heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p


def add_para(text, bold=False, italic=False, color=None, size=10, indent=False, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_shaded_para(lines, bg_hex="F8F9FA", border_hex="DEE2E6"):
    """Add a shaded box using a 1-cell table"""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    # Set background color
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg_hex)
    tcPr.append(shd)
    # Add lines
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        bold = line.startswith("**") and line.endswith("**")
        is_bullet = line.startswith("* ")
        text = line.lstrip("* ").strip("*")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        if is_bullet:
            p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.bold = bold
    return table


def add_criteria_table(rows):
    """Add evaluation criteria table: #, 통과조건(코드/설명), 통과조건(캡처)"""
    header = ["#", "통과 조건 (코드/설명)", "통과 조건 (캡처/증빙)"]
    table = doc.add_table(rows=len(rows) + 1, cols=3)
    table.style = "Table Grid"
    # Header
    for j, h in enumerate(header):
        cell = table.cell(0, j)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8.5)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "E9ECEF")
        tcPr.append(shd)
    # Data
    widths = [Cm(1.5), Cm(8.5), Cm(8.5)]
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i + 1, j)
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(8.5)
    # Column widths
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]
    return table


def add_submit_table(items):
    """Add submission items table"""
    table = doc.add_table(rows=len(items) + 1, cols=2)
    table.style = "Table Grid"
    # Header
    for j, h in enumerate(["제출물", "경로"]):
        cell = table.cell(0, j)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8.5)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "E9ECEF")
        tcPr.append(shd)
    for i, (title, path) in enumerate(items):
        table.cell(i + 1, 0).paragraphs[0].add_run(title).font.size = Pt(8.5)
        table.cell(i + 1, 1).paragraphs[0].add_run(path).font.size = Pt(8)
    widths = [Cm(8), Cm(10.5)]
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            cell.width = widths[j]
    return table


def add_zero_box(extra=None):
    lines = [
        "**0점 조건 (X-1, X-2 중 하나라도 미충족 시 해당 문제 전체 0점)**",
        "* 캡처/로그/실행결과 증빙 없음",
        "* 수정 전/후 비교 없음",
        "* 가설 2개 미만 (가설 요구 문제)",
        "* 최종 원인 명확히 특정 못함",
        "* 캡처에 date && hostname 출력 미포함",
    ]
    if extra:
        lines.append(f"* {extra}")
    return add_shaded_para(lines, bg_hex="FFF3F3", border_hex="FFCCCC")


def page_break():
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_heading("10주차 실전 장애 대응 시험", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

sub = doc.add_paragraph("Docker · Kafka · Spark · Airflow · Databricks 데이터 파이프라인")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub.runs:
    run.font.size = Pt(11)

add_hr()

# 시험 구조
add_heading("시험 구조", level=2)
add_shaded_para([
    "**총 10문제 (100점 만점) | 문제당 10점 | 총 20문제 (문제당 2개 세부문제)**",
    "* 각 문제는 X-1, X-2 두 개의 세부문제로 구성됩니다.",
    "* X-1, X-2 중 하나라도 틀리면 해당 문제 전체 0점 (이진 채점)",
    "* 채점 방식: 자동채점 60% + 구조화 수동채점 40%",
], bg_hex="EEF6FF", border_hex="B7D7FF")

doc.add_paragraph()

# 과제 제출 안내
add_heading("과제 제출 안내", level=2)
add_shaded_para([
    "아래 경우 모두 0점 처리",
    "* .docx / .hwp 등 문서 파일만 단독 제출",
    "* 코드를 이미지(캡처/사진)로 제출",
    "반드시 아래 폴더 구조를 지켜 ZIP 파일로 압축하여 제출하세요.",
], bg_hex="FFF3F3", border_hex="FFCCCC")
add_para("파일명: 데엔몇기_이름_10주차.zip", bold=True)
add_para("예시: 데엔3기_홍길동_10주차.zip")

doc.add_page_break()

# 증빙 원칙
add_heading("증빙 원칙 (AI 차단 설계)", level=3)
add_shaded_para([
    "* 모든 문제는 실행 환경 의존 — 동일 코드라도 환경에 따라 결과가 다름",
    "* 모든 문제는 가설 → 검증 → 증명 구조 강제",
    "* 중간 결과(로그, UI 캡처, diff) 제출 필수 — 최종 결과만 제출 시 0점",
    "* 캡처 시 반드시 date && hostname 명령 출력 포함",
], bg_hex="FFF8E1", border_hex="FFE082")

doc.add_paragraph()

# 제출물 공통 구조
add_heading("제출물 공통 구조 (모든 문제 동일)", level=3)
add_shaded_para([
    "1. 문제 재현 결과 (실패 상태 캡처)",
    "2. 실패 로그 (관련 로그 라인 인용)",
    "3. 원인 가설 2개 이상",
    "4. 실제 원인 특정 + 근거",
    "5. 수정 내용 (코드 포함)",
    "6. 수정 후 결과 (성공 상태 캡처)",
    "7. 수정 전/후 비교",
], bg_hex="F4FBF4", border_hex="B8DDB8")

page_break()


# ════════════════════════════════════════════════════════════════
def write_problem(num, title_text, theme, level_text, situation_lines,
                  sub1_title, sub1_reqs, sub1_criteria, sub1_code_submit,
                  sub2_title, sub2_reqs, sub2_criteria, sub2_code_submit,
                  submit_items, zero_extra=None):

    add_heading(f"문제 {num}: {title_text} [10점]", level=1)
    add_para(f"주제: {theme} | 난이도: {level_text}", italic=True)
    add_hr()

    add_heading("상황", level=2)
    add_shaded_para(situation_lines)
    doc.add_paragraph()

    # 문제 X-1
    add_heading(f"문제 {num}-1: {sub1_title}", level=2)
    add_heading("요구사항", level=3)
    for req in sub1_reqs:
        add_bullet(req)
    add_heading("평가 기준", level=3)
    add_criteria_table(sub1_criteria)
    if sub1_code_submit:
        doc.add_paragraph()
        add_heading("코드 제출물", level=3)
        for item in sub1_code_submit:
            add_bullet(item)
    doc.add_paragraph()

    # 문제 X-2
    add_heading(f"문제 {num}-2: {sub2_title}", level=2)
    add_heading("요구사항", level=3)
    for req in sub2_reqs:
        add_bullet(req)
    add_heading("평가 기준", level=3)
    add_criteria_table(sub2_criteria)
    if sub2_code_submit:
        doc.add_paragraph()
        add_heading("코드 제출물", level=3)
        for item in sub2_code_submit:
            add_bullet(item)
    doc.add_paragraph()

    add_heading("제출물 목록", level=2)
    add_submit_table(submit_items)
    doc.add_paragraph()
    add_zero_box(zero_extra)
    page_break()


# ════════════════════════════════════════════════════════════════
# PROBLEM 1
write_problem(
    num=1,
    title_text="보이지 않는 환경 — Docker 서비스 실패 진단",
    theme="환경 디버깅",
    level_text="기본",
    situation_lines=[
        "제공된 docker-compose.yaml을 실행하면 일부 서비스가 실패합니다.",
        "* Kafka broker가 기동 후 수초 내 crash",
        "* Spark executor가 Kafka에 연결 실패",
        "* 에러 메시지가 혼재되어 어떤 서비스가 root cause인지 불명확",
    ],
    sub1_title="서비스 실패 원인 분석",
    sub1_reqs=[
        "docker-compose up 실행 후 실패 로그 캡처",
        "docker ps / docker logs 결과 캡처",
        "원인 가설 최소 2개 제시 (각 가설에 대한 검증 방법 포함)",
        "실제 원인 1개 특정 + 로그 라인 인용 근거",
        "다른 가설 배제 이유 설명",
    ],
    sub1_criteria=[
        ["1-1a", "원인 가설 2개 이상 제시\n각 가설에 검증 방법 포함", "docker logs에서 실패 로그 라인 캡처\n+ date && hostname 포함"],
        ["1-1b", "실제 원인 특정\n+ 다른 가설 배제 논리", "로그 라인 인용 +\n원인-결과 인과관계 설명"],
    ],
    sub1_code_submit=None,
    sub2_title="서비스 복구 및 정상 동작 증명",
    sub2_reqs=[
        "수정 전/후 설정 파일 diff 제출",
        "docker-compose up 성공 캡처",
        "서비스 포트 오픈 여부 확인 (curl / nc 등)",
        "Kafka / Spark 연결 테스트 결과 캡처",
        "컨테이너 healthcheck 결과 캡처",
    ],
    sub2_criteria=[
        ["1-2a", "수정 전/후 설정 비교\n(diff 형태 제출)", "docker ps에서 모든 서비스\nUp 상태 캡처"],
        ["1-2b", "포트 오픈 + 연결 테스트\n명령어 제출", "Kafka produce/consume +\nSpark 연결 성공 캡처"],
    ],
    sub2_code_submit=["수정된 docker-compose.yaml (problem-1/answer/docker-compose-fixed.yaml)"],
    submit_items=[
        ("3개 서비스 로그 캡처 (date && hostname 포함)", "problem-1/answer/cascade_logs.png"),
        ("3개 서비스 로그 합본 (시간순 텍스트)", "problem-1/answer/cascade_logs.txt"),
        ("원인 분석 문서 (인과관계 포함)", "problem-1/answer/root_cause_analysis.md"),
        ("수정 전/후 설정 diff", "problem-1/answer/config_diff.txt"),
        ("수정 후 전체 서비스 healthy 캡처", "problem-1/answer/services_healthy.png"),
        ("Spark job 3회 성공 로그 캡처", "problem-1/answer/spark_job_success.png"),
    ],
    zero_extra="실행 성공만 제출하고 원인 설명이 없는 경우",
)

# ════════════════════════════════════════════════════════════════
# PROBLEM 2
write_problem(
    num=2,
    title_text="네트워크 미로 — Kafka Listener 및 네트워크 설정 디버깅",
    theme="환경 디버깅 (심화)",
    level_text="기본",
    situation_lines=[
        "모든 컨테이너가 Up 상태이지만, 외부 클라이언트에서 Kafka에 메시지를 발행/소비할 수 없습니다.",
        "* docker ps 상 모든 서비스 healthy",
        "* 컨테이너 내부에서는 Kafka 통신 정상",
        "* 호스트에서 localhost:9092 연결 시 timeout",
    ],
    sub1_title="네트워크 문제 진단",
    sub1_reqs=[
        "컨테이너 내부에서 kafka-console-producer 성공 캡처",
        "호스트에서 동일 명령 실패 캡처",
        "kafka broker 설정 중 listener 관련 설정 캡처",
        "원인 가설 2개 이상 (네트워크, listener, port mapping 등)",
        "실제 원인 특정 + 설정값 인용",
    ],
    sub1_criteria=[
        ["2-1a", "내부/외부 연결 차이 분석\n+ 가설 2개 이상", "내부 성공 + 외부 실패\n캡처가 동시에 보임"],
        ["2-1b", "ADVERTISED_LISTENERS\n설정 문제 특정", "broker 설정에서\n문제 설정값 인용"],
    ],
    sub1_code_submit=None,
    sub2_title="네트워크 설정 수정 및 E2E 통신 증명",
    sub2_reqs=[
        "수정 전/후 설정 diff",
        "호스트에서 produce → consume 성공 캡처",
        "Spark 컨테이너에서 Kafka 연결 성공 캡처",
        "전체 네트워크 흐름도 — Excalidraw(https://excalidraw.com/)로 작성 후 PNG 캡처 제출 (host → docker network → kafka → consumer 흐름 포함)",
    ],
    sub2_criteria=[
        ["2-2a", "listener 설정 수정 전/후 diff", "호스트에서 produce +\nconsume 성공 캡처"],
        ["2-2b", "네트워크 흐름도\n(Excalidraw PNG)", "Spark → Kafka\n연결 성공 캡처"],
    ],
    sub2_code_submit=["수정된 Kafka 설정 파일 (problem-2/answer/kafka-fixed.yaml 또는 server.properties)"],
    submit_items=[
        ("내부/외부 연결 비교 캡처", "problem-2/answer/internal_vs_external.png"),
        ("원인 분석 문서", "problem-2/answer/network_analysis.md"),
        ("수정된 Kafka 설정 파일", "problem-2/answer/kafka-fixed.yaml"),
        ("설정 diff", "problem-2/answer/listener_diff.txt"),
        ("E2E 통신 성공 캡처", "problem-2/answer/e2e_communication.png"),
    ],
    zero_extra="로그 근거 없이 설정만 변경한 경우",
)

# ════════════════════════════════════════════════════════════════
# PROBLEM 3
write_problem(
    num=3,
    title_text="시공간의 미아 — DAG 멱등성 검증",
    theme="데이터 정합성",
    level_text="중급",
    situation_lines=[
        "제공된 DAG를 2회 이상 실행하면 데이터가 중복 적재됩니다.",
        "* late data 존재 (이벤트 시간 vs 처리 시간 차이)",
        "* timezone 혼재 (UTC vs KST)",
        "* 동일 데이터가 반복 적재됨",
    ],
    sub1_title="중복 적재 현상 재현 및 진단",
    sub1_reqs=[
        "1차 실행 후 SELECT count(*) 결과 캡처",
        "2차 실행 후 SELECT count(*) 결과 캡처 (증가 확인)",
        "중복 검증 쿼리 (GROUP BY + HAVING count > 1) 결과 캡처",
        "중복 발생 원인 가설 2개 이상",
        "실제 원인 특정 (timezone / dedup 기준 / write mode 등)",
    ],
    sub1_criteria=[
        ["3-1a", "중복 검증 쿼리 제출\n(GROUP BY + HAVING)", "1차/2차 실행 후\nrow count 비교 캡처 (수치 증가)"],
        ["3-1b", "중복 원인 가설 2개 이상\n+ 실제 원인 특정", "duplicate 쿼리 결과에서\n중복 row 존재 캡처"],
    ],
    sub1_code_submit=["중복 검증 쿼리 파일 (problem-3/answer/duplicate_check.sql)"],
    sub2_title="멱등성 보장 구현 및 검증",
    sub2_reqs=[
        "dedup 전략 설명 (upsert / delete-insert / staging-merge 등)",
        "event time 기준 처리 구현 (processing time 아닌 event time)",
        "수정된 DAG/코드 제출",
        "수정 후 2회 실행 결과 checksum 비교",
        "수정 후 row count 동일 여부 캡처",
        "수정 후 duplicate 검증 쿼리 결과 (0건) 캡처",
    ],
    sub2_criteria=[
        ["3-2a", "dedup 전략 코드 구현\n+ event time 기준 처리", "2회 실행 후\nrow count 동일 캡처"],
        ["3-2b", "idempotent 설계 논리 문서", "duplicate 검증 = 0건\n+ checksum 동일 캡처"],
    ],
    sub2_code_submit=[
        "수정된 DAG 파일 (problem-3/answer/idempotent_dag.py)",
        "중복 검증 쿼리 (problem-3/answer/duplicate_check.sql)",
    ],
    submit_items=[
        ("1차/2차 실행 row count 비교 캡처", "problem-3/answer/row_count_comparison.png"),
        ("수정 전 중복 검증 결과 캡처", "problem-3/answer/duplicates_before.png"),
        ("수정된 DAG 파일", "problem-3/answer/idempotent_dag.py"),
        ("중복 검증 쿼리", "problem-3/answer/duplicate_check.sql"),
        ("dedup 전략 설명 문서", "problem-3/answer/dedup_strategy.md"),
        ("수정 후 멱등성 검증 캡처 (row count + checksum)", "problem-3/answer/idempotency_proof.png"),
    ],
    zero_extra="diff 없이 '동일하다'고만 주장하는 경우",
)

# ════════════════════════════════════════════════════════════════
# PROBLEM 4
write_problem(
    num=4,
    title_text="시간의 함정 — Timezone 혼재 데이터 정합성",
    theme="데이터 정합성 (심화)",
    level_text="중급",
    situation_lines=[
        "파이프라인에 UTC와 KST가 혼재되어 있어, 날짜 경계 부근 데이터가 잘못된 파티션에 적재됩니다.",
        "* Producer는 UTC 기준 timestamp 발행",
        "* Spark job은 KST 기준으로 날짜 파티션 생성",
        "* 2025-04-01 00:00~08:59 UTC 데이터가 3월 31일/4월 1일에 분산",
    ],
    sub1_title="Timezone 불일치 현상 재현",
    sub1_reqs=[
        "UTC 2025-04-01 00:00 ~ 08:59 범위 데이터 100건 이상 발행",
        "Spark job 실행 후 날짜별 파티션 row count 조회",
        "기대값 (4월 1일에 100건) vs 실제값 (3월 31일/4월 1일 분산) 비교",
        "원인 가설 2개 이상",
        "실제 원인 특정 (코드에서 timezone 처리 지점 인용)",
    ],
    sub1_criteria=[
        ["4-1a", "날짜 경계 데이터 발행\n코드/명령어 제출", "파티션별 row count에서\n데이터 분산 확인 캡처"],
        ["4-1b", "timezone 불일치 원인 특정\n+ 코드 지점 인용", "기대값 vs 실제값\n수치 비교 캡처"],
    ],
    sub1_code_submit=["경계 데이터 발행 스크립트 (problem-4/answer/boundary_producer.py)"],
    sub2_title="Timezone 통일 및 정합성 복구",
    sub2_reqs=[
        "수정 전/후 코드 diff",
        "수정 후 동일 데이터 재처리",
        "수정 후 파티션별 row count = 기대값 캡처",
        "수정 후 2회 실행 시 결과 동일 (멱등성 유지) 캡처",
    ],
    sub2_criteria=[
        ["4-2a", "timezone 통일 코드 수정 (diff)", "수정 후 파티션별\nrow count = 기대값 캡처"],
        ["4-2b", "— (자동채점)", "2회 실행 후\nrow count 동일 캡처"],
    ],
    sub2_code_submit=[
        "수정된 Spark job (problem-4/answer/spark_job_fixed.py)",
        "경계 데이터 발행 스크립트 (problem-4/answer/boundary_producer.py)",
    ],
    submit_items=[
        ("수정 전 파티션 분포 캡처", "problem-4/answer/partition_before.png"),
        ("원인 분석 문서", "problem-4/answer/timezone_analysis.md"),
        ("수정된 Spark job", "problem-4/answer/spark_job_fixed.py"),
        ("수정 전/후 코드 diff", "problem-4/answer/code_diff.txt"),
        ("수정 후 파티션 분포 + 멱등성 검증 캡처", "problem-4/answer/partition_after.png"),
    ],
    zero_extra="timezone 관련 코드 지점을 특정하지 못한 경우",
)

# ════════════════════════════════════════════════════════════════
# PROBLEM 5
write_problem(
    num=5,
    title_text="보이지 않는 실패 — SUCCESS인데 데이터 없음",
    theme="디버깅 / Observability",
    level_text="고급",
    situation_lines=[
        "파이프라인의 모든 단계가 SUCCESS를 반환하지만, 최종 DB에 데이터가 없습니다.",
        "* Airflow DAG: SUCCESS",
        "* Spark job: SUCCESS (exit code 0)",
        "* PostgreSQL target 테이블: 0 rows",
    ],
    sub1_title="Silent Failure 원인 추적",
    sub1_reqs=[
        "Airflow SUCCESS 화면 캡처",
        "Spark job SUCCESS 로그 캡처",
        "PostgreSQL SELECT count(*) = 0 캡처",
        "데이터 흐름 각 단계 확인: Kafka offset → Spark read count → Spark write count → DB count",
        "데이터가 사라지는 지점 특정 + 로그 라인 인용",
        "원인 가설 2개 이상",
    ],
    sub1_criteria=[
        ["5-1a", "데이터 흐름 4단계\n각각의 count 추적", "Kafka(N) → Spark read(N)\n→ Spark write(?) → DB(0) 수치 캡처"],
        ["5-1b", "데이터 유실 지점 특정\n+ 로그 라인 인용", "가설 2개 이상\n+ 실제 원인 근거"],
    ],
    sub1_code_submit=None,
    sub2_title="수정 및 데이터 적재 증명",
    sub2_reqs=[
        "수정 전/후 코드 diff",
        "수정 후 파이프라인 실행 결과 캡처",
        "DB target 테이블 row 존재 확인 캡처",
        "staging vs final 테이블 count 일치 확인",
    ],
    sub2_criteria=[
        ["5-2a", "수정 코드 diff 제출", "실행 후 DB row > 0 캡처"],
        ["5-2b", "— (자동채점)", "staging count = final count 캡처"],
    ],
    sub2_code_submit=["수정된 파이프라인 코드 (problem-5/answer/)"],
    submit_items=[
        ("SUCCESS 상태 캡처 (Airflow + Spark)", "problem-5/answer/success_but_empty.png"),
        ("데이터 흐름 추적 캡처", "problem-5/answer/data_flow_trace.png"),
        ("원인 분석 문서", "problem-5/answer/silent_failure_analysis.md"),
        ("수정된 코드", "problem-5/answer/"),
        ("수정 전/후 코드 diff", "problem-5/answer/code_diff.txt"),
        ("수정 후 DB 적재 증명 캡처", "problem-5/answer/data_loaded.png"),
    ],
    zero_extra="'~인 것 같다' 수준의 추정만 있고 로그 인용이 없는 경우",
)

# ════════════════════════════════════════════════════════════════
# PROBLEM 6
write_problem(
    num=6,
    title_text="로그의 거짓말 — 불완전 로그 속 진짜 원인 찾기",
    theme="디버깅 / Observability (심화)",
    level_text="고급",
    situation_lines=[
        "Spark job이 간헐적으로 실패합니다. 로그에는 여러 에러가 섞여 있습니다.",
        "* 10회 실행 중 3~4회 실패",
        "* 에러 로그: ConnectionTimeout, OOM, TaskNotSerializable 혼재",
        "* 실패 시점마다 에러 메시지가 다름",
        "* 일부 로그는 의도적으로 누락/오도용으로 삽입됨 (로그 트랩)",
    ],
    sub1_title="로그 분류 및 진짜 원인 식별",
    sub1_reqs=[
        "제공된 로그 파일에서 에러 라인 전체 목록 추출",
        "각 에러를 '관련(relevant)' / '무관(irrelevant)'으로 분류 + 분류 근거",
        "무관 로그를 배제한 이유 설명",
        "원인 가설 2개 이상",
        "실제 root cause 특정 + 로그 시계열 분석 (어떤 에러가 먼저 발생했는지)",
    ],
    sub1_criteria=[
        ["6-1a", "에러 분류표 제출\n(relevant / irrelevant + 각 분류 근거)", "로그 파일에서\n에러 라인 추출 캡처"],
        ["6-1b", "root cause 특정\n+ 시계열 분석", "가설 2개 이상\n+ 배제 논리"],
    ],
    sub1_code_submit=None,
    sub2_title="수정 및 안정성 증명",
    sub2_reqs=[
        "수정 전/후 코드 diff",
        "수정 후 5회 연속 실행 결과 캡처 (각 실행 시각 포함)",
        "5회 모두 SUCCESS 확인",
        "수정이 root cause를 해결하는 이유 설명",
    ],
    sub2_criteria=[
        ["6-2a", "수정 코드 diff\n+ 수정 이유 설명", "5회 실행 결과 캡처 (시각 포함)"],
        ["6-2b", "— (자동채점)", "5회 모두 SUCCESS (실패 0건)"],
    ],
    sub2_code_submit=["수정된 Spark 코드 (problem-6/answer/spark_job_fixed.py)"],
    submit_items=[
        ("에러 분류표", "problem-6/answer/error_classification.md"),
        ("원인 분석 문서", "problem-6/answer/root_cause_analysis.md"),
        ("수정된 Spark 코드", "problem-6/answer/spark_job_fixed.py"),
        ("수정 전/후 코드 diff", "problem-6/answer/code_diff.txt"),
        ("5회 연속 실행 성공 캡처", "problem-6/answer/five_runs_success.png"),
    ],
    zero_extra="로그 분류 없이 바로 수정만 제출한 경우",
)

# ════════════════════════════════════════════════════════════════
# PROBLEM 7
write_problem(
    num=7,
    title_text="리소스 절벽 — Spark OOM 분석",
    theme="Spark 성능",
    level_text="심화",
    situation_lines=[
        "제공된 Spark job이 OOM(OutOfMemory)으로 실패합니다.",
        "* executor에서 'java.lang.OutOfMemoryError: Java heap space' 발생",
        "* 일부 executor에서만 OOM (skew 의심)",
        "* shuffle read 과다 발생",
        "* Spark UI에서 stage 3 task retry 다수 확인",
    ],
    sub1_title="OOM 원인 분석",
    sub1_reqs=[
        "실패 로그 캡처 (OOM 에러 라인 포함)",
        "Spark UI Jobs 탭 캡처 (실패 stage 표시)",
        "Spark UI Stage Details 캡처 (task별 shuffle read/write, GC time)",
        "executor별 메모리 사용량 비교 (skew 확인)",
        "원인 가설 2개 이상 (skew / shuffle partition 과다 / memory config 부족 등)",
        "실제 원인 특정 + Spark UI 수치 인용",
    ],
    sub1_criteria=[
        ["7-1a", "Spark UI 해석\n(stage, task, shuffle 수치)", "OOM 에러 로그 +\nSpark UI 캡처 (Jobs/Stage Details)"],
        ["7-1b", "skew/shuffle 원인 분석\n+ 수치 인용", "가설 2개 이상\n+ executor별 비교"],
    ],
    sub1_code_submit=None,
    sub2_title="성능 튜닝 및 성공 실행",
    sub2_reqs=[
        "변경한 Spark config 목록 + 각 변경 이유",
        "튜닝 후 성공 실행 결과 캡처",
        "Spark UI 전/후 비교 (Duration, Shuffle, GC time)",
        "결과 row count 정상 확인",
    ],
    sub2_criteria=[
        ["7-2a", "config 변경 목록\n+ 각 변경 이유", "튜닝 후 성공 실행\n+ row count 정상"],
        ["7-2b", "— (자동채점: Duration 비교)", "Spark UI 전/후\nDuration 감소 캡처"],
    ],
    sub2_code_submit=["튜닝 config 적용된 실행 스크립트 (problem-7/answer/run_tuned.sh 또는 spark_job_tuned.py)"],
    submit_items=[
        ("실패 로그 + Spark UI 캡처 (전)", "problem-7/answer/spark_ui_before.png"),
        ("원인 분석 문서", "problem-7/answer/oom_analysis.md"),
        ("튜닝 config + 이유 문서", "problem-7/answer/tuning_config.md"),
        ("튜닝된 실행 스크립트", "problem-7/answer/run_tuned.sh"),
        ("성공 실행 + Spark UI 캡처 (후)", "problem-7/answer/spark_ui_after.png"),
        ("전/후 성능 비교표 캡처", "problem-7/answer/performance_comparison.png"),
    ],
    zero_extra="Spark UI 캡처 없이 코드만 제출한 경우",
)

# ════════════════════════════════════════════════════════════════
# PROBLEM 8
write_problem(
    num=8,
    title_text="클라우드 탈출 — Databricks 성능 최적화",
    theme="Spark 성능 + Databricks",
    level_text="심화",
    situation_lines=[
        "문제 7에서 로컬 환경의 리소스 한계를 확인했습니다.",
        "이제 Databricks 환경에서 동일 job을 실행하고, 클러스터 설정을 최적화합니다.",
        "* 로컬에서 OOM 났던 job을 Databricks에서 실행",
        "* 기본 cluster 설정으로는 비효율적 (과도한 리소스 또는 부족한 리소스)",
    ],
    sub1_title="Databricks Cluster 구성 및 기본 실행",
    sub1_reqs=[
        "Databricks cluster 생성 (worker type, node 수, Spark version 명시)",
        "cluster 설정 캡처",
        "notebook에 코드 업로드 후 실행",
        "실행 결과 + row count 확인 캡처",
        "로컬 실행 vs Databricks 실행 시간 비교",
    ],
    sub1_criteria=[
        ["8-1a", "cluster 설정\n(type, nodes, version) 명시", "Databricks cluster\n설정 화면 캡처"],
        ["8-1b", "notebook 실행 성공\n+ 결과 row count", "실행 결과 캡처\n+ 로컬 대비 시간 비교"],
    ],
    sub1_code_submit=["Databricks notebook 파일 (problem-8/answer/notebook.ipynb 또는 .py)"],
    sub2_title="클러스터 튜닝 및 전/후 비교",
    sub2_reqs=[
        "튜닝 전 실행 시간 기록",
        "cluster config 변경 (shuffle partitions, memory, executor 수 등)",
        "변경 이유 설명 (각 설정이 왜 성능에 영향을 주는지)",
        "튜닝 후 실행 시간 기록",
        "Spark UI (Databricks) 전/후 캡처",
        "성능 개선율 (%) 계산",
    ],
    sub2_criteria=[
        ["8-2a", "config 변경 목록\n+ 각 변경 이유", "Spark UI (Databricks)\n전/후 캡처"],
        ["8-2b", "성능 개선율 계산 (수치 기반)", "실행 시간 전/후\n비교표 캡처"],
    ],
    sub2_code_submit=["튜닝된 notebook (problem-8/answer/notebook_tuned.ipynb 또는 .py)"],
    submit_items=[
        ("Databricks cluster 설정 캡처", "problem-8/answer/cluster_config.png"),
        ("notebook 파일", "problem-8/answer/notebook.py"),
        ("notebook 실행 결과 캡처", "problem-8/answer/notebook_result.png"),
        ("튜닝된 notebook", "problem-8/answer/notebook_tuned.py"),
        ("튜닝 전/후 Spark UI 캡처", "problem-8/answer/databricks_spark_ui.png"),
        ("성능 비교 문서 (변경 이유 + 개선율)", "problem-8/answer/performance_tuning.md"),
    ],
    zero_extra="tuning 이유 설명 없이 config만 변경한 경우",
)

# ════════════════════════════════════════════════════════════════
# PROBLEM 9
write_problem(
    num=9,
    title_text="데이터 신뢰 붕괴 — KPI 이상 탐지",
    theme="데이터 품질 / 분석",
    level_text="최고급",
    situation_lines=[
        "대시보드의 KPI 수치가 비정상입니다. 데이터는 정상 적재되어 있습니다.",
        "* 일별 매출이 갑자기 2배로 증가",
        "* 사용자 수가 전일 대비 50% 감소",
        "* 데이터 적재 건수는 정상 범위",
        "* KPI 정의는 제공되지 않음 — 직접 정의해야 함",
    ],
    sub1_title="KPI 정의 및 이상 탐지",
    sub1_reqs=[
        "KPI 정의 문서 (매출, 사용자 수 등 최소 3개 KPI 정의 + 계산 수식)",
        "이상 탐지 쿼리/코드 제출",
        "정상 기간 vs 이상 기간 수치 비교표",
        "이상 원인 가설 2개 이상 (join fan-out, duplicate, schema 변경 등)",
        "각 가설에 대한 검증 방법 제시",
    ],
    sub1_criteria=[
        ["9-1a", "KPI 3개 이상 정의\n+ 계산 수식", "정상 vs 이상 기간\n수치 비교표 캡처"],
        ["9-1b", "이상 탐지 쿼리/코드\n+ 가설 2개 이상", "이상 수치를\n정량적으로 증명하는 캡처"],
    ],
    sub1_code_submit=["이상 탐지 쿼리 (problem-9/answer/anomaly_detection.sql)"],
    sub2_title="원인 특정 및 KPI 정상화",
    sub2_reqs=[
        "실제 원인 특정 (fan-out / duplicate / 데이터 왜곡 등)",
        "원인 특정 근거 (쿼리 결과, 데이터 샘플 등)",
        "수정 코드/쿼리 제출",
        "수정 후 KPI 계산 결과 = 정상 범위 캡처",
        "data lineage 설명 (원본 → 변환 → 최종 테이블 경로)",
    ],
    sub2_criteria=[
        ["9-2a", "실제 원인 특정\n+ 근거 데이터", "수정 전/후\nKPI 값 비교 캡처"],
        ["9-2b", "data lineage 설명\n(원본 → 최종)", "수정 후 KPI =\n정상 범위 캡처"],
    ],
    sub2_code_submit=[
        "수정 쿼리/코드 (problem-9/answer/fix_queries.sql)",
        "이상 탐지 쿼리 (problem-9/answer/anomaly_detection.sql)",
    ],
    submit_items=[
        ("KPI 정의 문서", "problem-9/answer/kpi_definition.md"),
        ("이상 탐지 쿼리", "problem-9/answer/anomaly_detection.sql"),
        ("원인 분석 문서", "problem-9/answer/root_cause_analysis.md"),
        ("수정 쿼리/코드", "problem-9/answer/fix_queries.sql"),
        ("수정 전/후 KPI 비교 캡처", "problem-9/answer/kpi_comparison.png"),
        ("data lineage 설명", "problem-9/answer/data_lineage.md"),
    ],
    zero_extra="KPI 정의 없이 이상만 보고하는 경우",
)

# ════════════════════════════════════════════════════════════════
# PROBLEM 10
write_problem(
    num=10,
    title_text="Lineage 지옥 — 데이터 왜곡 추적 및 복구",
    theme="데이터 품질 / Lineage (심화)",
    level_text="최고급",
    situation_lines=[
        "문제 9에서 KPI 이상을 발견한 후 추가 조사 결과, 데이터 왜곡이 여러 단계에 걸쳐 발생합니다.",
        "* staging 테이블에는 정상 데이터",
        "* final 테이블에서 매출 2배 (join fan-out 의심)",
        "* 특정 날짜에만 사용자 수 50% 감소 (필터 조건 오류 의심)",
        "* 데이터 트랩: timestamp 1초 차이, key 중복, schema 변경이 숨어 있음",
    ],
    sub1_title="다단계 데이터 왜곡 추적",
    sub1_reqs=[
        "각 단계별 row count / 합계 비교표 (staging vs transform vs final)",
        "fan-out 검증 쿼리: join 전후 row count 비교",
        "duplicate 검증 쿼리: GROUP BY + HAVING count > 1",
        "timestamp 1초 차이로 인한 이중 집계 발견 쿼리",
        "schema 변경 이력 확인 (ALTER TABLE / 컬럼 추가 시점)",
        "원인 가설 3개 이상 (fan-out, duplicate, filter 오류, schema drift 등)",
        "각 왜곡 지점별 실제 원인 특정",
    ],
    sub1_criteria=[
        ["10-1a", "3단계 count 비교표\n+ fan-out/duplicate 검증 쿼리", "staging vs final\n수치 불일치 캡처"],
        ["10-1b", "왜곡 지점 3개 이상 식별\n+ 각각의 원인 특정", "데이터 트랩 발견\n(timestamp/key/schema) 근거 캡처"],
    ],
    sub1_code_submit=["왜곡 추적 쿼리 모음 (problem-10/answer/trace_queries.sql)"],
    sub2_title="전체 파이프라인 복구 및 Lineage 문서화",
    sub2_reqs=[
        "각 왜곡 지점별 수정 코드/쿼리",
        "수정 후 각 단계 count 일관성 확인",
        "수정 후 KPI 정상화 확인 (문제 9 KPI 기준)",
        "전체 data lineage 다이어그램 (원본 → kafka → staging → transform → final)",
        "각 변환 단계에서의 필터/집계/조인 로직 설명",
        "향후 재발 방지를 위한 data quality check 제안 (최소 3개)",
    ],
    sub2_criteria=[
        ["10-2a", "수정 코드/쿼리\n+ 각 단계 count 일관성", "수정 후 staging = transform\n= final count 캡처"],
        ["10-2b", "lineage 다이어그램\n+ quality check 제안 3개", "수정 후 KPI =\n정상 범위 캡처"],
    ],
    sub2_code_submit=[
        "수정 쿼리 (problem-10/answer/fix_queries.sql)",
        "왜곡 추적 쿼리 (problem-10/answer/trace_queries.sql)",
    ],
    submit_items=[
        ("다단계 count 비교표 캡처", "problem-10/answer/stage_count_comparison.png"),
        ("왜곡 추적 분석 문서", "problem-10/answer/distortion_analysis.md"),
        ("왜곡 추적 쿼리", "problem-10/answer/trace_queries.sql"),
        ("수정 쿼리/코드", "problem-10/answer/fix_queries.sql"),
        ("수정 후 정합성 검증 캡처", "problem-10/answer/consistency_proof.png"),
        ("data lineage 다이어그램", "problem-10/answer/data_lineage_diagram.png"),
        ("data quality check 제안서", "problem-10/answer/quality_check_proposal.md"),
    ],
    zero_extra="이상을 수치로 증명하지 못하거나 원인 특정이 없는 경우",
)


# ════════════════════════════════════════════════════════════════
# APPENDIX — 채점 기준 총괄표
add_heading("부록: 채점 기준 총괄표", level=1)
add_hr()

add_heading("점수 구조", level=2)
summary_table = doc.add_table(rows=11, cols=5)
summary_table.style = "Table Grid"
headers = ["문제", "주제", "난이도", "배점", "세부문제"]
for j, h in enumerate(headers):
    cell = summary_table.cell(0, j)
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.bold = True
    run.font.size = Pt(8.5)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "E9ECEF")
    tcPr.append(shd)

rows_data = [
    ["1", "환경 디버깅 (기본)", "기본", "10", "1-1 원인 분석 + 1-2 복구 증명"],
    ["2", "환경 디버깅 (심화)", "기본", "10", "2-1 네트워크 진단 + 2-2 E2E 통신"],
    ["3", "멱등성 (DAG)", "중급", "10", "3-1 중복 재현 + 3-2 멱등성 구현"],
    ["4", "멱등성 (Timezone)", "중급", "10", "4-1 TZ 불일치 재현 + 4-2 TZ 통일"],
    ["5", "디버깅 (Silent Fail)", "고급", "10", "5-1 유실 추적 + 5-2 적재 증명"],
    ["6", "디버깅 (로그 트랩)", "고급", "10", "6-1 로그 분류 + 6-2 안정성 증명"],
    ["7", "Spark 성능 (OOM)", "심화", "10", "7-1 OOM 분석 + 7-2 튜닝 성공"],
    ["8", "Spark + Databricks", "심화", "10", "8-1 클러스터 실행 + 8-2 튜닝 비교"],
    ["9", "데이터 품질 (KPI)", "최고급", "10", "9-1 KPI 정의+탐지 + 9-2 정상화"],
    ["10", "품질 + Lineage", "최고급", "10", "10-1 왜곡 추적 + 10-2 복구+Lineage"],
]
for i, row in enumerate(rows_data):
    for j, val in enumerate(row):
        cell = summary_table.cell(i + 1, j)
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(8.5)

doc.add_paragraph()

# 채점 프로세스
add_heading("채점 프로세스", level=2)

add_heading("1단계: 자동채점 (60%)", level=3)
add_shaded_para([
    "* 코드 실행 성공 여부",
    "* 결과값 검증 (row count, checksum, KPI 계산)",
    "* duplicate 체크",
    "* 성능 threshold 만족 여부",
    "**자동채점 50% 미만 시 실격 처리**",
], bg_hex="EEF6FF", border_hex="B7D7FF")

doc.add_paragraph()

add_heading("2단계: 구조화 수동채점 (40%)", level=3)
add_shaded_para([
    "* 로그 분석 능력 (특정 로그 라인 명시 + 의미 설명)",
    "* 가설 설정 능력 (가설 2개 이상 + 각 가설 근거)",
    "* 원인 특정 능력 (root cause 명확 + 다른 가설 배제)",
    "* 수정 근거 (수정 내용 명확 + 왜 효과 있는지 설명)",
    "**채점자 지침: 설명을 읽고 이해하려 하지 말고, 체크리스트 항목이 있는지만 확인하라**",
], bg_hex="F4FBF4", border_hex="B8DDB8")

doc.add_paragraph()
add_zero_box()

doc.save(OUTPUT)
print(f"DOCX generated: {OUTPUT}")
